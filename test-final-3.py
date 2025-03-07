import streamlit as st
import pdfplumber
import json
import os
import shutil
import openai
from docx import Document
import re

# Nh·∫≠p OpenAI API Key
api_key = st.text_input("üîë Nh·∫≠p OpenAI API Key:", type="password")
if not api_key:
    st.warning("Vui l√≤ng nh·∫≠p OpenAI API Key ƒë·ªÉ s·ª≠ d·ª•ng ·ª©ng d·ª•ng.")
    st.stop()

# C·∫•u h√¨nh OpenAI client
client = openai.OpenAI(api_key=api_key)
def clean_text(text):
    """L√†m s·∫°ch vƒÉn b·∫£n ƒë·ªÉ lo·∫°i b·ªè k√Ω t·ª± NULL v√† c√°c k√Ω t·ª± ƒëi·ªÅu khi·ªÉn kh√¥ng h·ª£p l·ªá."""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)  # Lo·∫°i b·ªè k√Ω t·ª± ƒëi·ªÅu khi·ªÉn
    return text.strip()
# H√†m tr√≠ch xu·∫•t vƒÉn b·∫£n t·ª´ PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += clean_text(page_text) + "\n"
    except Exception as e:
        st.error(f"L·ªói khi x·ª≠ l√Ω PDF: {e}")
        return None
    return text
def clean_json_response(response_text):
    """L√†m s·∫°ch ph·∫£n h·ªìi GPT ƒë·ªÉ lo·∫°i b·ªè c√°c k√Ω t·ª± kh√¥ng mong mu·ªën tr∆∞·ªõc khi ph√¢n t√≠ch JSON."""
    response_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', response_text)
    response_text = response_text.strip()
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    return response_text.strip()
# H√†m g·ªçi GPT ƒë·ªÉ tr√≠ch xu·∫•t th√¥ng tin
def extract_info_with_gpt(text):
    prompt = f"""
    Tr√≠ch xu·∫•t th√¥ng tin t·ª´ vƒÉn b·∫£n CV v√† tr·∫£ v·ªÅ JSON h·ª£p l·ªá:
    {{
        "Name": "",
        "Email": "",
        "Phone": "",
        "Skills": [],  
        "Experience": [],  
        "Education": [],  
        "Certifications": [],  
        "Languages": [],  
        "Strengths": [],  
        "Weaknesses": [],  
        "Additional information": []
    }}
    For the **Languages** field, include:
    - The candidate's native language based on their nationality (e.g., Vietnamese for a candidate from Vietnam).
    - Any foreign language certifications (e.g., TOEIC score) and the corresponding language proficiency level (e.g., English with a proficiency level based on the score).

    For **Strengths and Weaknesses**, analyze the candidate's work experience to identify:
    - **Strengths:** Key skills and attributes demonstrated through their experience.
    - **Weaknesses:** Areas for improvement or challenges faced in their roles.
    CV text:
    {text}
    """

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an expert in extracting information from CVs (resumes) and images with 10 years of experience in getting the exact information needed to recruit suitable positions for the company."

            "**Context:** I will provide you with resumes of candidates (which can be one or more) or image files containing text."

            "**Your task** is to extract information from the resumes and images I provide (I have taken the text from the resume, and the image will be provided to you below) and return the output as a JSON file."

            "Some of the most important information required for each candidate includes:"
            "- Name"
            "- Email"
            "- Phone number"
            "- Skills"
            "- Experience (including: position, timeline, responsibilities)"
            "- Education (including: degree, institution, timeline, GPA)"
            "- Certifications"
            "- Languages (including proficiency based on nationality and language certifications)"
            "- Strengths (based on the candidate's experience and job description)"
            "- Weaknesses (based on the candidate's experience and job description)"
            "- Additional information (including identification and visa details if provided)"

            "**Task:** Extract the following information from the CV text and return it as JSON."

            "**Output:** JSON file format"

            "***Note:** I can provide you with the text, but in that text will be a synthesis of many resumes of different candidates.*"

            "**REMEMBER:** The output should only be in JSON format." },
            {"role": "user", "content": prompt}
        ]
    )
    # return json.loads(response.choices[0].message.content.strip())
    extracted_text = response.choices[0].message.content.strip()
    cleaned_text = clean_json_response(extracted_text)
    # return json.loads(response.choices[0].message.content.strip())
    return json.loads(cleaned_text)
# L∆∞u JSON
def save_to_json(data_list, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data_list, f, ensure_ascii=False, indent=4)

# T·∫°o file Word
def create_word_file(text_list, output_path):
    try:
        doc = Document()
        for idx, (filename, text) in enumerate(text_list):
            doc.add_heading(f"File {idx + 1}: {filename}", level=1)
            doc.add_paragraph(clean_text(text))
            doc.add_page_break()
        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"L·ªói khi t·∫°o file Word: {e}")
        return False

# ·ª®ng d·ª•ng Streamlit
st.title("üìÑ PDF to JSON & AI Assistant")
st.write("üîπ T·∫£i l√™n CV (PDF), tr√≠ch xu·∫•t n·ªôi dung v√† ƒë·∫∑t c√¢u h·ªèi v·ªõi tr·ª£ l√Ω ·∫£o!")

uploaded_files = st.file_uploader("üì§ Ch·ªçn file PDF", type=["pdf"], accept_multiple_files=True)
temp_dir = "temp"
os.makedirs(temp_dir, exist_ok=True)

if uploaded_files:
    extracted_data = {}
    extracted_texts = {}

    for uploaded_file in uploaded_files:
        temp_file_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        text = extract_text_from_pdf(temp_file_path)
        if text:
            extracted_texts[uploaded_file.name] = text
            extracted_info = extract_info_with_gpt(text)
            extracted_info["Filename"] = uploaded_file.name
            extracted_data[uploaded_file.name] = extracted_info
        os.remove(temp_file_path)

    # T·∫°o file Word v√† JSON
    word_output = os.path.join(temp_dir, "extracted_texts.docx")
    if create_word_file(list(extracted_texts.items()), word_output):
        with open(word_output, "rb") as file:
            st.download_button("üì• T·∫£i file Word", file, file_name="extracted_texts.docx")

    json_output = os.path.join(temp_dir, "extracted_data.json")
    save_to_json(list(extracted_data.values()), json_output)
    with open(json_output, "rb") as file:
        st.download_button("üì• T·∫£i file JSON", file, file_name="extracted_data.json")

    # Chatbot v·ªõi tr·ª£ l√Ω ·∫£o
    st.subheader("üí¨ Chat v·ªõi tr·ª£ l√Ω ·∫£o v·ªÅ n·ªôi dung CV")

    if len(uploaded_files) > 1:
        selected_cv = st.selectbox("Ch·ªçn CV ƒë·ªÉ t∆∞∆°ng t√°c:", list(extracted_texts.keys()))
    else:
        selected_cv = list(extracted_texts.keys())[0]
    if "openai_model" not in st.session_state:
        st.session_state.openai_model = "gpt-3.5-turbo"
    if "messages" not in st.session_state:
        st.session_state.messages = []
    # Display chat messages from history on app rerun
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    if "openai_model" not in st.session_state:
        st.session_state["openai_model"] = "gpt-3.5-turbo"
    text_CV = f"D∆∞·ªõi ƒë√¢y l√† n·ªôi dung CV c·ªßa ·ª©ng vi√™n:\n{extracted_texts[selected_cv]}"
    prompt = st.chat_input("H·ªèi tr·ª£ l√Ω ·∫£o v·ªÅ CV n√†y:")
    if prompt:
         # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(prompt)
        with st.chat_message("assistant"):
            stream = client.chat.completions.create(
            model=st.session_state["openai_model"],
            messages=[{"role": "system", "content": "Play the role of a professional HR, with 10 years of experience in finding potential candidates suitable for the company based on the CV (resume) they send Context: I will provide you with information of each CV (resume) in text form, from which I will ask you some questions related to the CV (resume) of this candidate Task: Please provide the most accurate and closest information to the question I asked, helping me have the most objective view of this candidate so that I can decide whether to hire him or not Tone: solemn, dignified, straightforward, suitable for the office environment, recruitment. Below is the content of the candidate's CV"  + extracted_texts[selected_cv] },
                {"role": "user", "content": prompt}],
                stream=True,
            )
                # if hasattr(response.choices[0].delta, "content"):    
            response = st.write_stream(stream)
        st.session_state.messages.append({"role": "assistant", "content": response})
    shutil.rmtree(temp_dir, ignore_errors=True)
